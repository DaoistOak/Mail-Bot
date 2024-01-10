import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from docx import Document
import openpyxl
import os
import re

# Function to save the starting_row to a text file
def save_starting_row(file_path, starting_row):
    with open(file_path, 'w') as file:
        file.write(str(starting_row))

# Function to read the starting_row from a text file
def read_starting_row(file_path):
    if not os.path.exists(file_path):
        with open(file_path, 'w') as file:
            file.write('1')  # Start from 1 if the file doesn't exist
    
    with open(file_path, 'r') as file:
        return file.read().strip()

def process_excel_rows(file_path, start_row):
    start_row = int(start_row)  # Ensure start_row is an integer

    wb = openpyxl.load_workbook(file_path)
    sheet = wb.active

    rows_content = []  # Variable to store row content

    for row_num in range(start_row, start_row + 20):
        cell = sheet[f"A{row_num}"]
        rows_content.append(cell.value)

        cell.font = openpyxl.styles.Font(color='FF0000')  # Change text color to red

    start_row += 20

    wb.save(file_path)
    return start_row, rows_content

def read_subject(file_path):
    doc = Document(file_path)
    subject = ""
    for para in doc.paragraphs:
        subject += para.text + "\n"  # Add each paragraph's text to the subject string with a newline

    return subject

def read_docx(file_path):
    doc = Document(file_path)
    content = []
    for para in doc.paragraphs:
        formatted_para = []
        for run in para.runs:
            # Get run-level properties: text, bold, color, etc.
            text = run.text
            bold = run.bold if run.bold else None
            color = run.font.color.rgb if run.font.color else None
            font_name = run.font.name if run.font.name else 'Arial'
            font_size = run.font.size.pt if run.font.size else '10'
            underline = run.font.underline if run.font.underline else None
            alignment = para.alignment if para.alignment else 'center'
            # Store the text and its formatting properties in a tuple
            formatted_text = (text, bold, color, font_name, font_size, underline, alignment)
            formatted_para.append(formatted_text)
        # Store the formatted paragraph in a list
        content.append(formatted_para)
    return content

def convert_rgb_to_hex(rgb):
    if rgb is None:
        return None
    # Convert RGB tuple to hexadecimal representation (HTML color code)
    return '#{:02x}{:02x}{:02x}'.format(rgb[0], rgb[1], rgb[2])

def generate_email_content(content, doc_link):
    email_content = ""
    for para in content:
        for run in para:
            text, bold, color, font_name, font_size, underline, alignment = run
            
            hex_color = convert_rgb_to_hex(color)
            
            # Replace tab characters with non-breaking spaces for indentation
            text = text.replace('\t', '&nbsp;&nbsp;&nbsp;&nbsp;')
            
            font_style = f"font-family: {font_name};"
            if font_size:
                font_style += f" font-size: {font_size}pt;"
            if hex_color:
                font_style += f" color: {hex_color};"
            
            # Construct the HTML with font styles
            formatted_text = f"<span style='text-align: center; font-weight: {'bold' if bold else 'normal'}; text-decoration: {'underline' if underline else 'none'};{font_style}'>{text}</span>"
            
            email_content += formatted_text
        email_content += "<br>"  # Adding line breaks between paragraphs
        
    # Concatenate the doc_link variable at the end of the email content
    email_content += f'<a href="{doc_link}">{doc_link}</a>'
    return email_content

def send_email(sender_email, sender_password, receiver_email, subject, content):
    # Set up the SMTP server
    smtp_server = 'smtp.gmail.com'
    smtp_port = 587  # Gmail SMTP port

    # Create a secure connection to the Gmail SMTP server
    server = smtplib.SMTP(smtp_server, smtp_port)
    server.starttls()  # Enable TLS encryption

    # Login to your Gmail account
    server.login(sender_email, sender_password)

    # Create message container
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['Bcc'] = ','.join(receiver_email)
    msg['Subject'] = subject

    # Generate HTML content for the email
    print('Generating Content')
    email_content = generate_email_content(content, doc_link)

    # Attach HTML content to the email
    print('Attaching E-mail')
    msg.attach(MIMEText(email_content, 'html'))

    # Send email
    print('Sending Email From -', sender_email)
    print('......')
    print('Sending Email to -', receiver_email)
    server.sendmail(sender_email, receiver_email, msg.as_string())
    # Quit SMTP server
    server.quit()

# Input for sender's email and password
sender_email = input("Enter your email: ")
sender_password = input("Enter your app key: ")
# Loop counter
loop_counter = 1
# Loop to perform actions 23 times
for _ in range(23):
    row_file = 'starting_row.txt'
    starting_row = read_starting_row(row_file)
    file_path = 'Gmail Directory 2078.xlsx'
    starting_row, selected_emails = process_excel_rows(file_path, starting_row)

    # Print the content of the selected rows in the desired format
    #formatted_selected_emails = ', '.join(selected_emails)

    # Example usage:
    receiver_email = selected_emails
    subject = read_subject('HR Promo Subject 2080.docx')  # Replace with your subject
    content = read_docx('HR Promo Letter 2080.docx')  # Replace with your document path
    doc_link = ""  # Replace with your document link

    # Sending the email
    send_email(sender_email, sender_password, receiver_email, subject, content)
    save_starting_row(row_file, starting_row)
    
    print(loop_counter)
    loop_counter += 1
    
    print('Done, Next one!')
    print('==================')

